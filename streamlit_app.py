import streamlit as st
import streamlit.components.v1 as components
import io
import fitz  # PyMuPDF per modificare PDF
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes
from docx import Document

st.set_page_config(page_title="Editor & Comparator PDF/Word", layout="wide")

if "selected_words" not in st.session_state:
    st.session_state.selected_words = []
if "annotations" not in st.session_state:
    st.session_state.annotations = []
if "uploaded_file" not in st.session_state:
    st.session_state.uploaded_file = None

st.title("Modifica, Annota e Confronta Documenti")
st.markdown("Carica un documento PDF o Word, seleziona testo, evidenzia, aggiungi annotazioni e confronta i documenti.")

# **Caricamento del file**
uploaded_file = st.file_uploader("Carica un file (PDF o Word)", type=["pdf", "docx"])

if uploaded_file:
    st.session_state.uploaded_file = uploaded_file

if st.session_state.uploaded_file:
    file_name = st.session_state.uploaded_file.name
    file_type = file_name.split(".")[-1]

    col1, col2 = st.columns(2)

    if file_type == "pdf":
        with col1:
            st.subheader("Documento Originale")
            pdf_bytes = st.session_state.uploaded_file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")

            extracted_text = ""
            for page in convert_from_bytes(pdf_bytes):
                extracted_text += pytesseract.image_to_string(page, lang="eng") + "\n"

            if extracted_text.strip():
                st.session_state.selected_words = extracted_text.split()
                st.text_area("Testo Estratto (OCR)", extracted_text, height=300, disabled=True)
            else:
                components.html('<iframe src="https://mozilla.github.io/pdf.js/web/viewer.html" width="100%" height="600px"></iframe>', height=650)

        with col2:
            st.subheader("Modifica il Documento")
            modified_text = st.text_area("Modifica il testo", extracted_text, height=300)

            # **Evidenziazione del testo**
            highlight_word = st.text_input("Inserisci una parola da evidenziare")
            if st.button("Evidenzia Parola"):
                if highlight_word and highlight_word in modified_text:
                    st.session_state.selected_words.append(highlight_word)

            st.write("### Parole evidenziate:")
            for word in st.session_state.selected_words:
                st.markdown(f"- 🟡 {word}")

            # **Aggiunta di annotazioni personalizzate**
            annotation_text = st.text_area("Scrivi un'annotazione")
            annotation_page = st.number_input("Pagina annotazione", min_value=1, max_value=len(doc), step=1)
            if st.button("Aggiungi Annotazione"):
                st.session_state.annotations.append({"text": annotation_text, "page": annotation_page})

            st.write("### Annotazioni salvate:")
            for ann in st.session_state.annotations:
                st.markdown(f"- 📌 **Pagina {ann['page']}:** {ann['text']}")

            # **Salvataggio PDF modificato**
            if st.button("Salva PDF Modificato"):
                pdf_buffer = io.BytesIO()
                for page_num, page in enumerate(doc):
                    for word in st.session_state.selected_words:
                        text_instances = page.search_for(word)
                        for inst in text_instances:
                            page.add_highlight_annot(inst)

                    for ann in st.session_state.annotations:
                        if ann["page"] - 1 == page_num:
                            page.insert_text((50, 50), ann["text"], fontsize=12, color=(1, 0, 0))

                doc.save(pdf_buffer)
                pdf_buffer.seek(0)

                st.download_button(
                    label="Scarica PDF Modificato",
                    data=pdf_buffer,
                    file_name="documento_modificato.pdf",
                    mime="application/pdf"
                )

    elif file_type == "docx":
        with col1:
            st.subheader("Documento Originale")
            doc = Document(st.session_state.uploaded_file)
            original_text = "\n".join([p.text for p in doc.paragraphs])
            st.text_area("Testo originale", original_text, height=500, disabled=True)

        with col2:
            st.subheader("Modifica il Documento")
            modified_text = st.text_area("Modifica il testo", original_text, height=500)

            if st.button("Esporta documento modificato (Word)"):
                modified_doc = Document()
                for line in modified_text.split("\n"):
                    modified_doc.add_paragraph(line)

                doc_buffer = io.BytesIO()
                modified_doc.save(doc_buffer)
                doc_buffer.seek(0)

                st.download_button(
                    label="Scarica Documento Word",
                    data=doc_buffer,
                    file_name="documento_modificato.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    if st.button("Carica un nuovo documento"):
        st.session_state.uploaded_file = None
        st.session_state.selected_words = []
        st.session_state.annotations = []
        st.experimental_rerun()
